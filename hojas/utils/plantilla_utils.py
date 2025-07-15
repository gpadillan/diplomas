import os
import tempfile
import pandas as pd
from docx import Document
from docx.shared import Pt
import win32com.client as win32
import pythoncom

def limpiar(valor):
    if pd.isna(valor):
        return ""
    return str(valor).strip()

def reemplazar_campos_en_docx(doc: Document, campos: dict):
    for p in doc.paragraphs:
        for campo, valor in campos.items():
            if campo in p.text:
                for run in p.runs:
                    if campo in run.text:
                        run.text = run.text.replace(campo, valor)
                        if campo in ("{{NOMBRE}}", "{{APELLIDOS}}"):
                            run.font.size = Pt(37)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                reemplazar_campos_en_docx(cell, campos)

def convertir_docx_a_pdf(docx_path: str, output_pdf: str, password: str = "1234"):
    pythoncom.CoInitialize()
    word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = False

    doc = word.Documents.Open(os.path.abspath(docx_path))
    temp_pdf = tempfile.mktemp(suffix=".pdf")
    doc.ExportAsFixedFormat(temp_pdf, 17)
    doc.Close(False)
    word.Quit()

    os.replace(temp_pdf, output_pdf)
    return output_pdf

def generar_documento(alumno, plantilla_path: str, sufijo_tipo: str = "") -> str:
    doc = Document(plantilla_path)

    fecha = alumno.get("FECHA")
    fecha_exp = alumno.get("FECHA EXPEDICIÓN")

    fecha_str = ""
    if pd.notna(fecha):
        try:
            fecha_str = pd.to_datetime(fecha).strftime("%d/%m/%Y")
        except Exception:
            fecha_str = str(fecha)

    fecha_exp_str = ""
    if pd.notna(fecha_exp):
        try:
            fecha_exp_str = pd.to_datetime(fecha_exp).strftime("%d/%m/%Y")
        except Exception:
            fecha_exp_str = str(fecha_exp)

    campos = {
        "{{NOMBRE}}": limpiar(alumno.get("NOMBRE")),
        "{{APELLIDOS}}": limpiar(alumno.get("APELLIDOS")),
        "{{DNI}}": limpiar(alumno.get("DNI ALUMNO")),
        "{{TITULO}}": limpiar(alumno.get("NOMBRE CURSO EXACTO EN TITULO")),
        "{{PROMOCION}}": limpiar(alumno.get("PROMOCION EN LA QUE FINALIZA")),
        "{{FECHA EXPEDICIÓN}}": fecha_exp_str,
        "{{FECHA}}": fecha_str,
        "{{NºTITULO}}": limpiar(alumno.get("Nº TITULO")),
    }

    print("📦 Campos generados:")
    for k, v in campos.items():
        print(f"{k}: {v}")

    reemplazar_campos_en_docx(doc, campos)

    temp_docx = "temp_documento.docx"
    doc.save(temp_docx)

    nombre_dni = campos["{{DNI}}"] or "sin_dni"
    sufijo_tipo = sufijo_tipo.upper() if sufijo_tipo else "SIN_TIPO"

    output_pdf = f"TITULO_{sufijo_tipo}_{nombre_dni}.pdf"
    return convertir_docx_a_pdf(temp_docx, output_pdf)
