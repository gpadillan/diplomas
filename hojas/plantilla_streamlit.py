import os
import tempfile
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx2pdf import convert  # ✔️ Compatible localmente en Windows
# Nota: docx2pdf no funciona en Streamlit Cloud, solo local

def limpiar(valor):
    if pd.isna(valor):
        return ""
    return str(valor).strip()

def reemplazar_campos_en_docx(doc: Document, campos: dict):
    for p in doc.paragraphs:
        for campo, valor in campos.items():
            if campo in p.text:
                for run in p.runs:
                    if campo in run.text:
                        run.text = run.text.replace(campo, valor)
                        if campo in ("{{NOMBRE}}", "{{APELLIDOS}}"):
                            run.font.size = Pt(37)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                reemplazar_campos_en_docx(cell, campos)

def convertir_docx_a_pdf(docx_path: str, output_pdf: str):
    # Sólo funciona en Windows localmente
    try:
        convert(docx_path, output_pdf)
        return output_pdf
    except Exception as e:
        print(f"⚠️ No se pudo convertir a PDF: {e}")
        return None

def generar_documento(alumno, plantilla_path: str, prefijo: str = "TITULO") -> str:
    doc = Document(plantilla_path)

    fecha = alumno.get("FECHA")
    fecha_exp = alumno.get("FECHA EXPEDICIÓN")

    fecha_str = pd.to_datetime(fecha).strftime("%d/%m/%Y") if pd.notna(fecha) else ""
    fecha_exp_str = pd.to_datetime(fecha_exp).strftime("%d/%m/%Y") if pd.notna(fecha_exp) else ""

    campos = {
        "{{NOMBRE}}": limpiar(alumno.get("NOMBRE")),
        "{{APELLIDOS}}": limpiar(alumno.get("APELLIDOS")),
        "{{DNI}}": limpiar(alumno.get("DNI ALUMNO")),
        "{{TITULO}}": limpiar(alumno.get("NOMBRE CURSO EXACTO EN TITULO")),
        "{{PROMOCION}}": limpiar(alumno.get("PROMOCION EN LA QUE FINALIZA")),
        "{{FECHA EXPEDICIÓN}}": fecha_exp_str,
        "{{FECHA}}": fecha_str,
        "{{NºTITULO}}": limpiar(alumno.get("Nº TITULO")),
    }

    print("📦 Campos generados:")
    for k, v in campos.items():
        print(f"{k}: {v}")

    reemplazar_campos_en_docx(doc, campos)

    temp_docx = tempfile.mktemp(suffix=".docx")
    doc.save(temp_docx)

    output_pdf = f"{prefijo}_{campos['{{DNI}}'] or 'sin_dni'}.pdf"
    pdf_path = convertir_docx_a_pdf(temp_docx, output_pdf)

    return pdf_path or temp_docx  # Retornar PDF si fue posible, si no el .docx
